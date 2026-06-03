import json
import io
from docx import Document
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils import timezone
from datetime import timedelta
from django.db.models.functions import TruncDate
from django.db.models import Count
from .utils.youtube import extract_video_id, get_video_transcript
from .utils.scraper import is_youtube_url, get_article_text
from .utils.ai_processor import generate_study_materials, chat_with_document, translate_document, eli5_document
from .models import StudyMaterial, Profile
from django.db.models import Q

def update_streak(user):
    import datetime
    profile, created = Profile.objects.get_or_create(user=user)
    
    last_active = profile.last_active_date
    if isinstance(last_active, datetime.datetime):
        last_active = last_active.date()
        
    today = timezone.now().date()
    if last_active < today:
        if (today - last_active).days == 1:
            profile.current_streak += 1
        else:
            profile.current_streak = 1
        profile.last_active_date = today
        profile.save()
    return profile.current_streak

def index(request):
    """
    Renders the main frontend page.
    """
    streak = 0
    if request.user.is_authenticated:
        streak = update_streak(request.user)
    return render(request, 'study_app/index.html', {'streak': streak})

@csrf_exempt
@login_required
def generate_materials(request):
    """
    API endpoint that receives a YouTube URL and returns the generated
    study materials (Notes, Summary, Quiz) in JSON format.
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            url = data.get('url', '')
            
            if not url:
                return JsonResponse({'error': 'No URL provided.'}, status=400)
            
            is_yt = is_youtube_url(url)
            video_id = ""
            
            if is_yt:
                video_id = extract_video_id(url)
                if not video_id:
                    return JsonResponse({'error': 'Invalid YouTube URL.'}, status=400)
                    
                # Check if this user has already generated materials for this video
                from .models import StudyMaterial
                existing_material = StudyMaterial.objects.filter(user=request.user, video_id=video_id).first()
                if existing_material:
                    return JsonResponse({
                        'notes': existing_material.notes,
                        'summary': existing_material.summary,
                        'quiz': existing_material.quiz,
                        'flashcards': existing_material.flashcards,
                        'mindmap': existing_material.mindmap,
                        'tags': existing_material.tags or '',
                        'video_id': video_id,
                        'title': existing_material.title,
                        'material_id': existing_material.id,
                        'share_id': str(existing_material.share_id)
                    }, status=200)
                    
                # 1. Fetch transcript
                content_text = get_video_transcript(video_id)
                title = f"Video Notes ({video_id})"
            else:
                # Treat as article
                content_text = get_article_text(url)
                title = f"Article Notes"
            
            # 2. Generate materials using AI
            materials = generate_study_materials(content_text)
            
            # 3. Save to database
            new_material = StudyMaterial.objects.create(
                user=request.user,
                video_id=video_id, # Can be empty for articles
                title=title,
                notes=materials.get('notes', ''),
                summary=materials.get('summary', ''),
                quiz=materials.get('quiz', []),
                flashcards=materials.get('flashcards', []),
                mindmap=materials.get('mindmap', ''),
                tags=''
            )
            materials['material_id'] = new_material.id
            materials['tags'] = ''
            
            # 4. Include video ID for embedding
            if is_yt:
                materials['video_id'] = video_id
            else:
                materials['video_id'] = None
                
            materials['share_id'] = str(new_material.share_id)
            
            return JsonResponse(materials, status=200)
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
            
    return JsonResponse({'error': 'Invalid request method.'}, status=405)

def register(request):
    """
    Handles user registration.
    """
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('index')
    else:
        form = UserCreationForm()
    return render(request, 'study_app/register.html', {'form': form})

@login_required
def dashboard(request):
    """
    Shows a history of all generated study materials for the logged-in user.
    """
    streak = update_streak(request.user)
    materials = StudyMaterial.objects.filter(user=request.user)
    
    # Search functionality
    query = request.GET.get('q')
    if query:
        materials = materials.filter(
            Q(title__icontains=query) | Q(summary__icontains=query) | Q(tags__icontains=query)
        )
        
    # Order by favorite first, then date
    materials = materials.order_by('-is_favorite', '-created_at')
    
    # Spaced Repetition Logic (Due for Review)
    today = timezone.now().date()
    due_dates = [today - timedelta(days=1), today - timedelta(days=3), today - timedelta(days=7)]
    due_for_review = materials.filter(
        created_at__date__in=due_dates, 
        is_mastered=False
    ).order_by('-created_at')
    
    # Analytics data (last 7 days)
    last_7_days = timezone.now() - timedelta(days=7)
    analytics = StudyMaterial.objects.filter(user=request.user, created_at__gte=last_7_days) \
                                     .annotate(date=TruncDate('created_at')) \
                                     .values('date') \
                                     .annotate(count=Count('id')) \
                                     .order_by('date')
    
    dates = [(timezone.now().date() - timedelta(days=i)).strftime('%Y-%m-%d') for i in range(6, -1, -1)]
    counts_by_date = {item['date'].strftime('%Y-%m-%d'): item['count'] for item in analytics}
    chart_data = [counts_by_date.get(d, 0) for d in dates]
    
    # Feature 6: Weekly Study Goal
    notes_generated_this_week = sum(chart_data)
    weekly_goal = 5 # Default goal
    
    return render(request, 'study_app/dashboard.html', {
        'materials': materials,
        'due_for_review': due_for_review,
        'streak': streak,
        'query': query or '',
        'chart_labels': json.dumps(dates),
        'chart_data': json.dumps(chart_data),
        'notes_generated_this_week': notes_generated_this_week,
        'weekly_goal': weekly_goal
    })

@csrf_exempt
@login_required
def toggle_favorite(request, material_id):
    if request.method == 'POST':
        material = get_object_or_404(StudyMaterial, id=material_id, user=request.user)
        material.is_favorite = not material.is_favorite
        material.save()
        return JsonResponse({'is_favorite': material.is_favorite})
    return JsonResponse({'error': 'Invalid request method.'}, status=405)

@csrf_exempt
@login_required
def toggle_mastered(request, material_id):
    if request.method == 'POST':
        material = get_object_or_404(StudyMaterial, id=material_id, user=request.user)
        material.is_mastered = not material.is_mastered
        material.save()
        return JsonResponse({'is_mastered': material.is_mastered})
    return JsonResponse({'error': 'Invalid request method.'}, status=405)

@csrf_exempt
@login_required
def update_tags(request, material_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            tags = data.get('tags', '')
            material = get_object_or_404(StudyMaterial, id=material_id, user=request.user)
            material.tags = tags
            material.save()
            return JsonResponse({'status': 'success', 'tags': material.tags})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid request method.'}, status=405)

@csrf_exempt
@login_required
def document_chat(request, material_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            question = data.get('question', '')
            material = get_object_or_404(StudyMaterial, id=material_id, user=request.user)
            
            # Combine notes and summary for context
            context_text = f"Title: {material.title}\n\nSummary:\n{material.summary}\n\nNotes:\n{material.notes}"
            
            answer = chat_with_document(context_text, question)
            return JsonResponse({'answer': answer})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid request method.'}, status=405)

@csrf_exempt
@login_required
def translate_content(request, material_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            target_language = data.get('language', 'Spanish')
            text_type = data.get('type', 'notes') # 'notes' or 'summary'
            
            material = get_object_or_404(StudyMaterial, id=material_id, user=request.user)
            text_to_translate = material.notes if text_type == 'notes' else material.summary
            
            translated_text = translate_document(text_to_translate, target_language)
            return JsonResponse({'translated_text': translated_text})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid request method.'}, status=405)

@csrf_exempt
@login_required
def eli5_content(request, material_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            text_type = data.get('type', 'notes')
            
            material = get_object_or_404(StudyMaterial, id=material_id, user=request.user)
            text_to_simplify = material.notes if text_type == 'notes' else material.summary
            
            eli5_text = eli5_document(text_to_simplify)
            return JsonResponse({'eli5_text': eli5_text})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid request method.'}, status=405)

@login_required
def export_word(request, material_id):
    from .models import StudyMaterial
    material = get_object_or_404(StudyMaterial, id=material_id, user=request.user)
    
    document = Document()
    document.add_heading(material.title, 0)
    
    document.add_heading('Summary', level=1)
    document.add_paragraph(material.summary)
    
    document.add_heading('Notes', level=1)
    # Basic HTML strip for notes (can be improved)
    import re
    notes_text = re.sub('<[^<]+>', '', material.notes)
    document.add_paragraph(notes_text)
    
    document.add_heading('Quiz', level=1)
    for i, q in enumerate(material.quiz):
        document.add_paragraph(f"Q{i+1}: {q.get('question')}")
        for opt in q.get('options', []):
            document.add_paragraph(f"- {opt}")
        document.add_paragraph(f"Answer: {q.get('answer')}\n")
    
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="StudyMaterial_{material_id}.docx"'
    return response

def shared_material(request, share_id):
    from .models import StudyMaterial
    material = get_object_or_404(StudyMaterial, share_id=share_id)
    # We can render a simplified template or reuse index.html with a flag
    return render(request, 'study_app/shared.html', {'material': material})

